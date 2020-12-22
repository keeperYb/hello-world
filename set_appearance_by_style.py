from docx import Document
from docx import document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE, WD_BUILTIN_STYLE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, Length, RGBColor
from docx.text.tabstops import TabStops
from docx.table import Table

"""表20 的raw_dict, 用于生成DocxTable obj"""
raw_dict_table_20 = {
    'table_name': '表20  服务器安全保密防护措施表',
    'column_names': [
        '序号',
        '服务器名称',
        '服务器IP',
        '部署位置',
        '“三合一”安装情况',
        '主机审计系统',
        '防病毒软件',
        '其他'],
    'row_list': [
        ['1', '服务器名称1', '服务器IP1', '部署位置1', '“三合一”安装情况1', '主机审计系统1', '防病毒软件1', '其他1'],
        ['2', '服务器名称2', '服务器IP2', '部署位置2', '“三合一”安装情况2', '主机审计系统2', '防病毒软件2', '其他2'],
    ]
}
test_docx_file = r'C:\Users\xuyb\Desktop\word\test_format.docx'
original_docx_file = r'C:\Users\xuyb\Desktop\word\表20.docx'
target_docx_file = r'C:\Users\xuyb\Desktop\word\表20 - 副本.docx'
"""
Paragraph styles in default template

    Normal
    Body Text
    Body Text 2
    Body Text 3
    Caption
    Heading 1
    Heading 2
    Heading 3
    Heading 4
    Heading 5
    Heading 6
    Heading 7
    Heading 8
    Heading 9
    Intense Quote
    List
    List 2
    List 3
    List Bullet
    List Bullet 2
    List Bullet 3
    List Continue
    List Continue 2
    List Continue 3
    List Number
    List Number 2
    List Number 3
    List Paragraph
    Macro Text
    No Spacing
    Quote
    Subtitle
    TOCHeading
    Title
"""


class DocxTable:
    """
    设定表格的类, 属性如下:
     - 表名
     - 列名 (as list)
     - 内容 (as rowList)
    """
    TABLE_NAME_STYLE = None  # store the style of table name(which is a para)
    TABLE_STYLE = None  # store the style(a Table obj)

    def __init__(self, raw_dict: dict):
        self.table_name = ''
        self.column_names = []
        self.row_list = [[]]

        for key in raw_dict.keys():
            setattr(self, key, raw_dict[key])
        pass

    pass


def set_paragraph_format():
    """
    set the format of one paragraph(for test, the paragraph is not saved)

    paragraph formats includes:
    # horizontal alignment
    # indentation
    # tab stops
    # paragraph spacing
    # line spacing
    # Pagination properties
    """
    # common setting of paragraph format
    dcmt = Document()
    assert isinstance(dcmt, document.Document)
    para = dcmt.add_paragraph()
    para_format = para.paragraph_format

    # the horizontal alignment
    para_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    alignment = para_format.alignment

    # the indentation, the first paragraph has unique 'first_line_indent',
    # which is interpreted relative to the left indent
    para_format.left_indent = Cm(5)  # set left indent to 5 cm
    left_indent = para_format.left_indent.cm

    # the tab stops
    tab_stops = para_format.tab_stops
    assert isinstance(tab_stops, TabStops)
    one_tab_stop = tab_stops.add_tab_stop(Cm(1.5))  # add one tab_stop of 1.5 cm
    tab_length_in_cm = one_tab_stop.position.cm
    the_first_tab_stop = tab_stops[0]  # use sequence semantics to indicate a tab_stop, equals to 'one_tab_stop'

    # paragraph spacing
    space_before = para_format.space_before
    space_after = para_format.space_after
    space_before = Pt(18)  # pt means Point, a measurement in printing(中文翻译: 磅, 为音译. 实际应为: 点)
    space_after = Pt(12)

    # line spacing
    line_spacing = para_format.line_spacing
    line_spacing_rule = para_format.line_spacing  # Exactly or Multiple times of para space

    # Pagination properties, all attributes are tri-state: True, False, None
    '''keep_together, keep one para in one page'''
    para_format.keep_together = True
    '''keep_with_next, keeps a paragraph on the same page as the subsequent paragraph '''
    para_format.keep_with_next = True
    '''page_break_before, causes a paragraph to be placed at the top of a new page'''
    para_format.page_break_before = True
    '''widow_control, breaks a page to avoid placing the first or last line of the paragraph on a separate page 
    from the rest of the paragraph.'''
    para_format.widow_control = True

    pass


def set_run_format():
    """set run(character) format
    run obj has lots of attributes , see Font API"""
    dcmt = Document()
    assert isinstance(dcmt, document.Document)
    run = dcmt.add_paragraph().add_run()
    run.text = 'test word format'
    font = run.font  # get the access of font obj

    # set font color
    font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    # font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1  # color blue, set by theme_color
    color_type = font.color.type  # the color type, typically is RGB or THEME

    run2 = dcmt.add_paragraph().add_run()
    run2.text = 'test word format 2'
    font2 = run2.font
    font2.color.rgb = RGBColor(207, 232, 203)  # 207, 232, 203 eye_protected_green
    dcmt.save(test_docx_file)
    pass


def handle_table_style():
    """get and set the table style"""
    dcmt = Document(original_docx_file)
    assert isinstance(dcmt, document.Document)
    for table in dcmt.tables:
        pass
    dcmt.save(original_docx_file)

    pass


def draw_table_in_docx(para_dict):
    """draw table to the target_docx_file"""
    docxTable_obj_20 = DocxTable(raw_dict_table_20)
    dcmt = Document(target_docx_file)
    assert isinstance(dcmt, document.Document)
    # 表名
    table_name_para = dcmt.add_paragraph(text=docxTable_obj_20.table_name)
    the_style = para_dict['table_name_style']
    styles = dcmt.styles
    # todo , what to do to determine a table's style???
    # styles.add_style('custom_table_style', builtin_styles[WD_BUILTIN_STYLE.TABLE_LIGHT_SHADING_ACCENT_1])
    try:
        the_style = styles.add_style(the_style.name, WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:  # already contains the style, then do nothing
        pass
    table_name_para.style = dcmt.styles[the_style.name]
    print('target 表20表名的样式名称: ' + table_name_para.style.name)

    # 表格正体
    table_20_rows = len(docxTable_obj_20.row_list) + 1  # rows + 1 , because of header
    table_20_cols = len(docxTable_obj_20.column_names)
    table_20 = dcmt.add_table(rows=table_20_rows, cols=table_20_cols)
    table_20.style = 'custom_table_style'
    ''' header'''
    header_cells = table_20.rows[0].cells
    for col in range(table_20_cols):
        header_cells[col].text = docxTable_obj_20.column_names[col]
        pass
    '''contents'''
    for content_row_index in range(1, table_20_rows):
        row_cells = table_20.row_cells(content_row_index)
        for cell in row_cells:
            cell_index = row_cells.index(cell)
            cell.text = docxTable_obj_20.row_list[content_row_index - 1][
                cell_index]  # -1 because row_list lack of header
            pass
        pass

    # finally set the appearance according to type...
    # todo now
    dcmt.save(target_docx_file)
    pass


def handle_table_style_simple():
    """set table style in a empty dcmt """
    demo_file_name = 'demo.docx'
    # dcmt = Document()
    # assert isinstance(dcmt, document.Document)
    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    # )
    #
    # table = dcmt.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc
    #     pass
    # # set table style
    # styles = dcmt.styles
    # table.style = 'LightShading-Accent2'
    # # table.style = styles[WD_BUILTIN_STYLE.TABLE_DARK_LIST]
    # dcmt.save('demo.docx')

    dcmt = Document(original_docx_file)
    assert isinstance(dcmt, document.Document)
    for table in dcmt.tables:
        table_style = table.style
        print(table_style.name)
        pass

    pass

def read_para_style():
    """read the style of one paragraph """
    dcmt = Document(original_docx_file)
    assert isinstance(dcmt, document.Document)
    for para in dcmt.paragraphs:
        if para.text == '表20  服务器安全保密防护措施表':
            the_para_style = para.style
            DocxTable.TABLE_NAME_STYLE = the_para_style
            # print('original 表20表名的样式名称: ' + the_para_style.name)
            pass
        pass
    pass


if __name__ == '__main__':
    # '''main area'''
    # read_para_style()
    # '''
    # the_para_dict has these keys:
    # # table_name_style
    # # table_style
    # # '''
    # handle_table_style()
    # the_para_dict = {
    #     'table_name_style': DocxTable.TABLE_NAME_STYLE,
    #     'table_style': DocxTable.TABLE_STYLE}
    # draw_table_in_docx(the_para_dict)
    # '''end main area'''
    # set_paragraph_format()
    # # set_run_format()
    # handle_table_style()

    handle_table_style_simple()
    print('Work Done!')
    pass
