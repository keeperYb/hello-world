"""测试 docxcompose"""
import os
import traceback
from typing import List

import docx
from docx import Document, document
from docxcompose.composer import Composer

# from database.database_helper import DataBaseHelper
from database import DataBaseHelper
from model.base import BaseNodeModel, BasePhysicalNodeModel
from utils.docx_compose_mailmerge.common_part import TemplateType, MyHeadingModel
from utils.docx_compose_mailmerge.concat_docxs import ParagraphStyleHelper
from utils.docx_compose_mailmerge.test_mailmerge import MailMergeHelper
from model.shared import DeReportNode, DeReportPhysicalNode, DeReportNetworkEnvironment, DeReportTopo, ReportAttachment
from utils.docx_helper import DocxHelper
from utils.io_helper import WorkDirHelper, OsHelper
from utils.snowflake import globalIdGenerator
from settings import TEMP_FOLDER_NAME

EVALUATION_DIR = os.path.join(WorkDirHelper.get_root_absolute_path(), r'template/docx')
DECLARE_DIR = os.path.join(WorkDirHelper.get_root_absolute_path(), r'template/docx_declare')
FENGPING_DECLARE_DIR = os.path.join(WorkDirHelper.get_root_absolute_path(), r'template/docx_fengping_declare')
FENGPING_DIR = os.path.join(WorkDirHelper.get_root_absolute_path(), r'template/docx_fengping')
YINGYONG_DIR = os.path.join(WorkDirHelper.get_root_absolute_path(), r'template/docx_yingYongXiTong')
YINGYONG_DECLARE_DIR = os.path.join(WorkDirHelper.get_root_absolute_path(), r'template/docx_yingYongXiTong_declare')

# TARGET_DIR = r'D:\desktop_files_and_folders\word\chapters'
EMPTY_TEMPLATE_DCMT_EVALUATION = os.path.join(EVALUATION_DIR, 'chapters/0_empty_base.docx')
EMPTY_TEMPLATE_DCMT_DECLARE = os.path.join(DECLARE_DIR, 'chapters/0_empty_base.docx')
EMPTY_TEMPLATE_DCMT_FENGPING = os.path.join(FENGPING_DIR, 'chapters/0_empty_base.docx')
EMPTY_TEMPLATE_DCMT_FENGPING_DECLARE = os.path.join(FENGPING_DECLARE_DIR, 'chapters/0_empty_base.docx')
# TARGET_DCMT = r'D:\desktop_files_and_folders\word\chapters\最终结果.docx'

DCMT_CENTRAL_NODE = os.path.join(EVALUATION_DIR, 'chapters/1st_part/1st_4_1_ZhongXinJieDian.docx')
DCMT_OTHER_NODE = os.path.join(EVALUATION_DIR, 'chapters/1st_part/1st_4_n_JieRuJieDian.docx')


class DocxComposer:
    def __init__(self, session, template_type: TemplateType, report_id, final_dcmt_path, random_postfix=''):
        """
        拼接docx文档的初始化阶段

        :param session:
        :param template_type:
        :param report_id:
        :param final_dcmt_path:
        :param random_postfix:
        """
        self.session = session
        self.template_type = template_type  # 模板的类型
        self.report_id = report_id
        self.final_dcmt_path = final_dcmt_path
        # todo 去de_declare找到所属的net_type, 判断 is_local(是否局域网)

        if template_type == TemplateType.CePing:
            self.empty_template = EMPTY_TEMPLATE_DCMT_EVALUATION
            self.parent_dir = EVALUATION_DIR
            pass
        elif template_type == TemplateType.ShenBao:
            self.empty_template = EMPTY_TEMPLATE_DCMT_DECLARE
            self.parent_dir = DECLARE_DIR
            pass
        elif template_type == TemplateType.FengPing:
            self.empty_template = EMPTY_TEMPLATE_DCMT_FENGPING
            self.parent_dir = FENGPING_DIR
            pass
        elif template_type == TemplateType.FengPingShenBao:
            self.empty_template = EMPTY_TEMPLATE_DCMT_FENGPING_DECLARE
            self.parent_dir = FENGPING_DECLARE_DIR
            pass
        else:  # 默认给的是EVALUATION的
            self.empty_template = EMPTY_TEMPLATE_DCMT_EVALUATION
            self.parent_dir = EVALUATION_DIR
            pass
        OsHelper.make_dir_if_not_existed(TEMP_FOLDER_NAME)
        self.temp_dir = os.path.join(TEMP_FOLDER_NAME, 'docx_compose_' + str(random_postfix))
        OsHelper.make_dir_if_not_existed(self.temp_dir)

        self.img_attached_ids = []  # 记录图片依附的记录 的id
        self.img_holders = []  # 图片在word文档中的占位符, 其中的例子形如: 图1:  xx
        pass

    def __compose_static_parts(self, first_part_document: str):
        """把完成的第一部分 和 固定的部分组合起来"""

        document_sequence_evaluation = ['封面',
                                        '报告签批页',
                                        '测评机构及委托方',
                                        '任务描述',
                                        first_part_document,
                                        '第二部分',
                                        '第三部分',
                                        '第四部分',
                                        '注意事项',
                                        '附件1',
                                        '附件2'
                                        ]
        document_sequence_declare = ['封面',
                                     '填写说明',
                                     '一_申请单位信息',
                                     first_part_document,
                                     '八_虚拟化安全防护与情况',
                                     '四_自查表',
                                     '自评估情况汇总表',
                                     ]
        document_sequence_fengping = document_sequence_evaluation
        document_sequence_fengping_declare = ['封面',
                                              '填写说明',
                                              '一_申请单位信息',
                                              first_part_document,
                                              '八_虚拟化安全防护与情况',
                                              '四_自查表',
                                              '自评估情况汇总表',
                                              ]
        if self.template_type == TemplateType.CePing:
            parent_dir = EVALUATION_DIR
            document_sequence = document_sequence_evaluation
            target_empty_dcmt = Document(EMPTY_TEMPLATE_DCMT_EVALUATION)
            pass
        elif self.template_type == TemplateType.ShenBao:
            parent_dir = DECLARE_DIR
            document_sequence = document_sequence_declare
            target_empty_dcmt = Document(EMPTY_TEMPLATE_DCMT_DECLARE)
            pass
        elif self.template_type == TemplateType.FengPing:
            parent_dir = FENGPING_DIR
            document_sequence = document_sequence_fengping
            target_empty_dcmt = Document(EMPTY_TEMPLATE_DCMT_FENGPING)
            pass
        elif self.template_type == TemplateType.FengPingShenBao:
            parent_dir = FENGPING_DECLARE_DIR
            document_sequence = document_sequence_fengping_declare
            target_empty_dcmt = Document(EMPTY_TEMPLATE_DCMT_FENGPING_DECLARE)
            pass
        else:  # 默认是测评
            parent_dir = EVALUATION_DIR
            document_sequence = document_sequence_evaluation
            target_empty_dcmt = Document(EMPTY_TEMPLATE_DCMT_EVALUATION)
            pass
        # 如果提出来公用, 需要如下的参数: empty_template_dcmt, docx_full_path_list, target_dcmt
        composer = Composer(target_empty_dcmt)
        self.__transfer_chinese_to_pinYin(document_sequence)
        for dcmt_simple_name in document_sequence:
            if dcmt_simple_name != first_part_document:  # 传来的first_part_document已经是全路径了
                static_file_path = os.path.join(parent_dir, 'chapters')
                dcmt_full_path = os.path.join(static_file_path, dcmt_simple_name + '.docx')
            else:
                dcmt_full_path = dcmt_simple_name
            try:
                composer.append(Document(dcmt_full_path))
            except:
                traceback.print_exc()
                print('异常的文档: ' + dcmt_simple_name)
            pass
        target_dcmt = os.path.join(self.temp_dir, '1_final_dcmt_to_mailmerge.docx')
        composer.save(target_dcmt)
        return target_dcmt
        pass

    def __auto_number_other_sections(self, one_full_path, temp_full_path, node_count, section_count, phy_node_count=''):
        """除了物理环境其他section的自动编号, 上面xx_count 都是str"""
        # todo auto_number
        dcmt = Document(one_full_path)
        assert isinstance(dcmt, document.Document)
        for pgh in dcmt.paragraphs:
            if pgh.style.name == ParagraphStyleHelper.MY_HEADING_4['name']:
                origin_text = pgh.text
                splitter = ' '
                try:
                    real_title = origin_text.split(splitter)[1]  # 标题
                except IndexError:  # 刚开始, 没有splitter来分...
                    real_title = origin_text
                    pass
                if phy_node_count:
                    prefix_auto_number = '3.%s.%s.%s%s' % (node_count, phy_node_count, section_count, splitter)
                else:
                    prefix_auto_number = '3.%s.%s%s' % (node_count, section_count, splitter)
                new_text = prefix_auto_number + real_title
                pgh.text = new_text
                pass
            dcmt.save(temp_full_path)
        pass

    def __auto_number_phy_node(self, one_full_path, temp_full_path, node_count, section_count, phy_node_count=''):
        """物理环境, 自动编号"""
        # todo auto_number
        dcmt = Document(one_full_path)
        assert isinstance(dcmt, document.Document)
        for pgh in dcmt.paragraphs:
            if pgh.style.name == ParagraphStyleHelper.MY_HEADING_4['name']:  # 物理环境一行, 是 MY_HEADING_4
                origin_text = pgh.text
                splitter = ' '
                try:
                    real_title = origin_text.split(splitter)[1]  # '物理环境'
                except IndexError:  # 刚开始, 没有splitter来分...
                    real_title = origin_text
                    pass
                if phy_node_count:
                    prefix_auto_number = '3.%s.%s%s' % (node_count, phy_node_count, splitter)
                else:  # 没有phy_node_count的情况
                    prefix_auto_number = '3.%s.%s%s' % (node_count, section_count, splitter)
                    pass
                new_text = prefix_auto_number + real_title
                pgh.text = new_text
                pass
            dcmt.save(temp_full_path)
            pass
        pass

    def __transfer_chinese_to_pinYin(self, chinese_list: list):
        """把原来的中文名的名称, 更换为对应拼音. 直接改变传入的chinese_list中的值"""
        from utils.docx_compose_mailmerge.common_part import chinese_pinYin_transform_dict
        transfer_dict = chinese_pinYin_transform_dict
        for i in range(len(chinese_list)):
            the_key = chinese_list[i]
            try:
                the_value = transfer_dict[the_key]
            except KeyError:
                continue  # 如果在dict中没找到chinese_key,这也是正常的,可以继续下一个循环
                pass
            chinese_list[i] = the_value  # 直接对原来的名称列表做修改
            pass
        pass

    def __get_phy_middle_part(self, phy_id):
        """组成物理环境的中间部分--图, 物理节点表, 物理安全措施表"""
        # 直接调用MailMergeHelper中的方法, 填写完成 'one_phy_id_middle_part', 然后存到temp_dir
        real_parent_dir = os.path.join(self.parent_dir, r'chapters/1st_part/sections')
        template_file = os.path.join(real_parent_dir, r'WuLiHuanJing_2_middle_contents.docx')
        target_dcmt = os.path.join(self.temp_dir, r'phy_middle_part_phyNode_%s.docx' % str(phy_id))
        mm_helper = MailMergeHelper()
        mm_helper.handle_phy_middle_part(session=self.session,
                                         phy_node_id=phy_id,
                                         template_file=template_file,
                                         target_file=target_dcmt, )
        one_phy_id_middle_part = target_dcmt
        return one_phy_id_middle_part
        pass

    def __generate_phy_template(self, node_obj: DeReportNode):
        """物理环境 小节需要根据物理节点的个数动态增加表格的数量
        并且需要返回mm过所有表格后的文档"""
        real_parent_dir = os.path.join(self.parent_dir, r'chapters/1st_part/sections')
        # 上部分
        template_up_part = os.path.join(real_parent_dir, r'WuLiHuanJing_1_up.docx')

        # 中部分
        node_id = node_obj.Id
        PhyNode = DeReportPhysicalNode
        phy_node_obj_list = self.session.query(PhyNode).filter(PhyNode.NodeId == node_id).all()
        if not phy_node_obj_list:
            return
        all_middle_parts = []  # including many one_middle_part
        for each_phy_obj in phy_node_obj_list:
            assert isinstance(each_phy_obj, DeReportPhysicalNode)
            phy_id = each_phy_obj.Id
            one_middle_part = self.__get_phy_middle_part(phy_id)
            # 其中有图片, 因此更新attached_id
            self.__add_img_attachment_ids_dynamic_parts(record_id=phy_id)
            all_middle_parts.append(one_middle_part)
            pass  # end for
        target_dcmt = os.path.join(self.temp_dir, r'phy_middle_part_node_%s.docx' % str(node_id))
        template_middle_part = self.compose_public(empty_template_dcmt=self.empty_template,
                                                   docx_full_path_sequence=all_middle_parts,
                                                   target_dcmt=target_dcmt)

        # 下部分
        template_down_part = os.path.join(real_parent_dir, r'WuLiHuanJing_3_down.docx')

        # 最后上, 中, 下合并
        phy_template_all_parts = [template_up_part, template_middle_part, template_down_part]
        target_dcmt = os.path.join(self.temp_dir, r'phy_all_part.docx')
        template_full_path = self.compose_public(empty_template_dcmt=self.empty_template,
                                                 docx_full_path_sequence=phy_template_all_parts,
                                                 target_dcmt=target_dcmt)

        return template_full_path
        pass

    def __compose_dynamic_sections(self, paras_dict: dict):
        """
        拼接物理节点下的小节

        :param paras_dict:  包含node_count: 用于小节的标号计数, node_obj: ReportNode obj
        :return: str, 表示18个小节拼成的临时文档的全路径
        """
        node_count = str(paras_dict['node_count'])
        # phy_node_count = str(paras_dict['phy_node_count'])
        node_obj = paras_dict['node_obj']
        section_count = 1  # 小节号, 从1开始
        section_name_sequence = [
            '物理环境',
            '网络环境',
            '系统定级及防护措施调整',
            '硬件平台',
            '软件平台',
            '安全保密设备',
            '应用系统',
            '身份鉴别',
            '安全审计',
            '电磁泄漏发射防护',
            '违规外联监控',
            '信息输入输出控制',
            '介质管控',
            '服务器安全保密防护情况',
            '安全保密管理机构与人员',
            '安全保密管理制度',
            '集成资质单位',
            '其他',
        ]
        self.__transfer_chinese_to_pinYin(section_name_sequence)  # 注意, section_name_sequence的中文被转为拼音..
        sections_parent_dir = os.path.join(self.parent_dir, r'chapters\1st_part\sections')
        dcmt_full_path_list = []
        for section_name in section_name_sequence:
            one_dcmt_name = section_name + '.docx'
            template_full_path = os.path.join(sections_parent_dir, one_dcmt_name)
            middle_full_path = os.path.join(self.temp_dir, one_dcmt_name)
            if section_name == 'WuLiHuanJing':
                template_full_path = self.__generate_phy_template(node_obj)
                self.__auto_number_phy_node(template_full_path, middle_full_path, node_count, str(section_count))
                pass
            else:
                if section_name == 'WangLuoHuanJing':
                    # 网络环境 中有图片需要加上record_id
                    cls = DeReportNetworkEnvironment
                    result = self.session.query(cls).filter(cls.NodeId == node_obj.Id).first()
                    if result:
                        network_env_id = result.Id
                    else:
                        network_env_id = None
                    self.__add_img_attachment_ids_dynamic_parts(record_id=network_env_id)
                    pass  # end if section_name ...
                self.__auto_number_other_sections(template_full_path, middle_full_path, node_count,
                                                  str(section_count))
                pass
            section_count += 1
            dcmt_full_path_list.append(middle_full_path)
            pass
        target_dcmt_simple_name = os.path.join(r'18section_to_mailmerge_%s.docx' % node_count)
        target_dcmt_full_path = os.path.join(self.temp_dir, target_dcmt_simple_name)
        # 得到第一部分的内容
        self.compose_public(empty_template_dcmt=self.empty_template,
                            docx_full_path_sequence=dcmt_full_path_list,
                            target_dcmt=target_dcmt_full_path)

        return target_dcmt_full_path
        pass

    def compose_DiYiBuFen(self, dynamic_part_dcmt: str):
        """
        拼接第一部分

        :param dynamic_part_dcmt: 已经拼接, mailmerge好的动态部分(节点)
        :return: 拼接好的第一部分的完整路径
        """
        first_part_parent_dir = os.path.join(self.parent_dir, r'chapters\1st_part')
        DiYiBuFen_sequence = [
            '第一部分1_总体情况',
            '第一部分2_总体网络拓扑图',
            '第一部分3_各节点情况',
        ]
        self.__transfer_chinese_to_pinYin(DiYiBuFen_sequence)
        full_path_DiYiBuFen_sequence = []  # 最终的full_path_list
        for each_dcmt in DiYiBuFen_sequence:
            real_each_dcmt = each_dcmt + '.docx'
            full_path_DiYiBuFen_sequence.append(os.path.join(first_part_parent_dir, real_each_dcmt))
            pass
        # 加上已经动态生成好的 节点 部分
        full_path_DiYiBuFen_sequence.append(dynamic_part_dcmt)
        target_dcmt = os.path.join(self.temp_dir, '1_1st_part_final_result.docx')
        self.compose_public(empty_template_dcmt=self.empty_template,
                            docx_full_path_sequence=full_path_DiYiBuFen_sequence,
                            target_dcmt=target_dcmt)
        return target_dcmt
        pass

    @staticmethod
    def compose_public(empty_template_dcmt, docx_full_path_sequence, target_dcmt):
        """合并docx的公用代码, 返回target_dcmt"""
        # 如果提出来公用, 需要如下的参数: empty_template_dcmt, docx_full_path_list, target_dcmt
        target_empty_dcmt = Document(empty_template_dcmt)
        composer = Composer(target_empty_dcmt)
        for dcmt_full_path in docx_full_path_sequence:
            try:
                composer.append(Document(dcmt_full_path))
            except Exception as e:
                traceback.print_exc()
                print(dcmt_full_path + ' runs into Exception ')
            pass
        composer.save(target_dcmt)
        return target_dcmt

    def __mailmerge_18sections(self, the_id, dcmt_path, paras_dict: dict):
        node_count = paras_dict['node_count']
        current_node_name = paras_dict['current_node_name']
        mailmerged_file_name = 'mailmerged_%s.docx' % (str(node_count))
        mailmerged_file_path = os.path.join(self.temp_dir, mailmerged_file_name)
        MailMergeHelper.handle_sections(session=self.session,
                                        phy_node_id=the_id,
                                        template_file=dcmt_path,
                                        target_file=mailmerged_file_path,
                                        current_node_name=current_node_name,
                                        )
        return mailmerged_file_path
        pass

    def __add_img_attachment_ids_dynamic_parts(self, record_id):
        """对于每个动态的小节, 把图片依附项的record_id加到self.img_attached_ids"""
        # # 物理环境图 依附的 de_report_physical_node 的 id
        # self.img_attached_ids.append(phy_node_id)
        #
        # # 节点拓扑图 依附的 de_report_network_environment 的 id
        # cls = DeReportNetworkEnvironment
        # if issubclass(cls, BaseNodeModel):
        #     result = self.session.query(cls).filter(cls.NodeId == phy_node_id).first()
        # elif issubclass(cls, BasePhysicalNodeModel):
        #     result = self.session.query(cls).filter(cls.PhysicalNodeId == phy_node_id).first()
        # else:  # 默认query NodeId
        #     result = self.session.query(cls).filter(cls.NodeId == phy_node_id).first()
        # if result:
        #     network_env_id = result.Id
        # else:
        #     network_env_id = None
        # self.img_attached_ids.append(network_env_id)
        # img_attach_ids = self.img_attached_ids
        self.img_attached_ids.append(record_id)
        pass

    def __get_one_node_document(self, node_obj: DeReportNode, node_count: int):
        """
        得到一个节点下的整个文档, 它包含多个phy_node的文档

        :param node_obj: ...
        :param node_count: 当前是第几个节点, 从1开始. 为1时, 节点为'中心节点'; 为2时, 节点为'接入节点1'; 以此类推
        :return: str: document_path
        """
        # 经过确认, 需要对18个小节做出区分处理:
        #  1. 物理环境 小节, 其中需要动态增添所有的物理节点
        #  2. 特殊的仍在使用phy_id的小节, 寻找它们对应的node_id作为外键
        #  3. 其他的小节都使用node_id作为外键

        dcmt_node_list = []
        # 加上小节编号, 一个lv3标题
        if node_count == 1:  # 第一个, 加上中心节点的标题
            tmp_dcmt = Document(DCMT_CENTRAL_NODE)
            assert isinstance(tmp_dcmt, document.Document)
            for pgh in tmp_dcmt.paragraphs:
                if '中心节点' in pgh.text:
                    pgh.text = '3.1' + '中心节点'
                pass
            dcmt_title_central_node = os.path.join(self.temp_dir, 'title_ZhongXinJieDian.docx')
            tmp_dcmt.save(dcmt_title_central_node)  # 更改了DCMT_OTHER_NODE 的实际内容
            dcmt_node_list.append(dcmt_title_central_node)  # todo, 如果 is_local, 就不append
        else:  # 其他, 加入接入节点n的标题
            tmp_dcmt = Document(DCMT_OTHER_NODE)
            assert isinstance(tmp_dcmt, document.Document)
            for pgh in tmp_dcmt.paragraphs:
                if '接入节点' in pgh.text:
                    pgh.text = '3.' + str(node_count) + '接入节点' + str(node_count - 1)
                pass
            dcmt_title_other_node = os.path.join(self.temp_dir, 'title_JieRuJieDian_%s.docx' % node_count)
            tmp_dcmt.save(dcmt_title_other_node)  # 更改了DCMT_OTHER_NODE 的实际内容
            dcmt_node_list.append(dcmt_title_other_node)   # todo, 如果 is_local, 就不append

        current_node_name = self.__get_current_node_name(node_count)
        paras_dict = {
            'node_obj': node_obj,
            'node_count': node_count,
            'current_node_name': current_node_name
        }
        dcmt_one_node = self.__compose_dynamic_sections(paras_dict)
        dcmt_one_node_mailmerged = self.__mailmerge_18sections(node_obj.Id, dcmt_one_node,
                                                               paras_dict)
        dcmt_node_list.append(dcmt_one_node_mailmerged)

        # phy_node_count = 1
        # for phy_node_obj in phy_node_obj_list:
        #     assert isinstance(phy_node_obj, orm_cls)
        #     phy_node_id = phy_node_obj.Id
        #     self.__add_img_attachment_ids__dynamic_parts(phy_node_id)
        #     dcmt_one_phy_node = self.__compose_dynamic_sections(paras_dict)
        #     # 接着mailmerge dcmt_one_phy_node
        #     current_node_name = self.__get_current_node_name(node_count)
        #     dcmt_one_node_mailmerged = self.__mailmerge_18sections(phy_node_id, dcmt_one_phy_node,
        #                                                                paras_dict, current_node_name)
        #
        #     dcmt_node_list.append(dcmt_one_node_mailmerged)
        #     phy_node_count += 1
        #     pass

        # 拼接所有的phy_node_dcmt, 成为一个dcmt_one_node
        dcmt_one_node = os.path.join(self.temp_dir, 'target_dcmt_node_%s.docx' % node_count)
        self.compose_public(empty_template_dcmt=self.empty_template,
                            docx_full_path_sequence=dcmt_node_list,
                            target_dcmt=dcmt_one_node)
        return dcmt_one_node

    def __get_current_node_name(self, node_count: int):
        """根据node_count 得到节点的名字"""
        if node_count == 1:
            current_node_name = '中心节点'
            pass
        else:
            current_node_name = '接入节点' + str(node_count - 1)
            pass
        return current_node_name
        pass

    def __get_all_nodes_document(self):
        """拼出包含所有node的文档(包含内容)"""
        session = self.session
        report_id = self.report_id
        # 得到所有的node_obj
        node_obj_list = session.query(DeReportNode).filter(DeReportNode.ReportId == report_id,
                                                           DeReportNode.RowStatus == 1).all()
        if not node_obj_list:
            return

        # 对每个node_obj, 得到对应的one_node_document, 并加到dcmt_one_node_list
        dcmt_node_list = []
        node_count = 1  # 节点计数器
        for node_obj in node_obj_list:
            dcmt_one_node = self.__get_one_node_document(node_obj, node_count)
            dcmt_node_list.append(dcmt_one_node)
            node_count += 1
            pass

        # 拼接dcmt_one_node_list中的所有document, 得到结果 dcmt_all_nodes
        dcmt_all_nodes = os.path.join(self.temp_dir, 'target_all_nodes.docx')
        self.compose_public(empty_template_dcmt=self.empty_template,
                            docx_full_path_sequence=dcmt_node_list,
                            target_dcmt=dcmt_all_nodes)
        return dcmt_all_nodes
        pass

    def __get_dynamic_part_document(self):
        """拼出第一部分的文档(包括内容), 返回结果文档的全路径"""
        # todo now 现在这里的变动更大了, 连总体网络拓扑图也会没有...
        dcmt_first_part = ''
        # 得到拼接了所有node的word文档
        document_all_nodes = self.__get_all_nodes_document()
        # 拼接第一部分的其他固定内容
        dcmt_first_part = self.compose_DiYiBuFen(document_all_nodes)
        return dcmt_first_part
        pass

    def __mailmerge_all(self, final_dcmt_to_mailmerge: str, final_target_dcmt: str):
        """
        完成所有固定部分的mailmerge

        :param final_dcmt_to_mailmerge: 将要进行最终mailmerge的模板
        :param final_target_dcmt: 最终生成的文档
        :return: None
        """
        MailMergeHelper.handle_static_parts(session=self.session,
                                            report_id=self.report_id,
                                            template_file=final_dcmt_to_mailmerge,
                                            target_file=final_target_dcmt,
                                            template_type=self.template_type)
        pass

    def __auto_number(self, final_target_dcmt: str):
        """最后为已经填上内容的文档 设置图片 和 表格的自动编号"""
        dcmt = Document(final_target_dcmt)
        global_table_count = 1
        global_image_count = 1
        if isinstance(dcmt, document.Document): pass
        for pgh in dcmt.paragraphs:
            if pgh.style.name == ParagraphStyleHelper.MY_TABLE_NUM['name']:
                origin_text = pgh.text
                postfix_table_name = origin_text.split('  ')[1]
                prefix_table_count = '表' + str(global_table_count) + '  '
                global_table_count += 1
                new_text = prefix_table_count + postfix_table_name
                pgh.text = new_text
                pass
            elif pgh.style.name == ParagraphStyleHelper.MY_IMAGE_NUM['name']:
                # change the image num
                origin_text = pgh.text
                postfix_image_name = origin_text.split('  ')[1]
                prefix_image_count = '图' + str(global_image_count) + '  '
                global_image_count += 1
                new_text = prefix_image_count + postfix_image_name
                self.img_holders.append(new_text)
                pgh.text = new_text
                pass
            else:
                pass
            pass
        dcmt.save(final_target_dcmt)
        pass

    def __add_img_attachment_ids_static(self):
        """
        对于固定部分, 加入依附的ids, 最后, self.img_attached_ids 中的id 需要根据word文档的图片排布顺序来排列
        :return:
        """
        # 总体网络拓扑图的Id
        cls = DeReportTopo
        result = self.session.query(cls).filter(cls.ReportId == self.report_id).first()
        if result:
            topo_id = result.Id
            pass
        else:
            topo_id = None
            pass
        tmp_list = [topo_id]
        tmp_list.extend(self.img_attached_ids)
        self.img_attached_ids.clear()  # clear 对 tmp_list有影响吗...
        self.img_attached_ids = tmp_list
        pass

    def __generate_img_dict_list(self):
        """生成img_dict_list, 返回img_dict_list
                每张图片有如下的结构:
        ---
        图x 图片名称     #这个称为holder
        [attached_id]  #这个称为attached_id, 记录了record_id
        ---
        """
        img_dir = os.path.join(WorkDirHelper.get_root_absolute_path(), 'image')
        attached_ids = self.img_attached_ids  # 根据其中有几个有效id, 加上对应个数的img_holder
        holders = self.img_holders
        img_dict_list = []
        idx = 0  # attached_ids, holders共用的index
        for each_record_id in attached_ids:
            one_img_dict = {}
            if not each_record_id:
                continue
            # 查询de_report_attachment
            cls = ReportAttachment
            result = self.session.query(cls).filter(cls.RecordId == each_record_id).first()
            if isinstance(result, cls): pass
            if result:
                simple_img_name = result.FileName
                full_path_img = os.path.join(img_dir, simple_img_name)
                one_img_dict['holder'] = holders[idx]  # 插入图片是根据顺序来的(idx)
                one_img_dict['imgPath'] = full_path_img
                img_dict_list.append(one_img_dict)
                pass
            else:
                pass
            idx += 1
            if idx >= len(holders):
                break
            pass  # end for each_record_id...
        return img_dict_list
        pass

    def __insert_images(self):
        """
        完成插入图片的工作. 在之前的步骤中, self.img_attached_ids已经加入了动态部分的依附id
        :return:
        """
        self.__add_img_attachment_ids_static()
        img_dict_list = self.__generate_img_dict_list()
        DocxHelper.insertPictures(self.final_dcmt_path, img_dict_list)
        pass

    def __add_page_break_at_end(self, dcmt_path: str):
        """add a page break at the end of the document"""
        dcmt = Document(dcmt_path)
        assert isinstance(dcmt, document.Document)
        dcmt.add_page_break()
        dcmt.save(dcmt_path)
        pass

    def final_compose(self):
        """
        最上层的mailmerge控制

        :return: str: final_target_dcmt
        """
        dynamic_part_document = self.__get_dynamic_part_document()
        self.__add_page_break_at_end(dynamic_part_document)
        final_document_to_mailmerge = self.__compose_static_parts(dynamic_part_document)
        final_target_dcmt = self.final_dcmt_path
        self.__mailmerge_all(final_document_to_mailmerge, final_target_dcmt)
        self.__auto_number(final_target_dcmt)
        self.__insert_images()
        self.generate_TOC(final_target_dcmt, final_target_dcmt)
        return final_target_dcmt

    pass

    @staticmethod
    def __add_toc(dcmt: document.Document, pgh, heading_list: List[MyHeadingModel]):
        """给dcmt 加上 TOC(table of content), 其实是替换字符串(很粗糙的列表, 连页码也没有)"""
        assert isinstance(pgh, docx.text.paragraph.Paragraph)
        heading_styles = [
            'my_heading_1',  # 0 tab
            'my_heading_2',  # 1 tab
            'my_heading_3',  # n-1 tabs ...
            'my_heading_4',
            'my_heading_5',
        ]
        total_toc = ''
        for each_heading_model in heading_list:
            lv_int = int(each_heading_model.lv[-1])  # 1,2...
            tabs = ''  # 在名称前加的制表符
            for i in range(lv_int - 1):
                tabs += '\t'
                pass
            one_line = tabs + each_heading_model.text
            total_toc += one_line + '\n'
            pass
        pgh.text = total_toc
        pass

    @staticmethod
    def generate_TOC(docx_full_path: str = '', target_file: str = ''):
        """为一个word文档制作Table Of Content(目录)"""
        if not docx_full_path:
            docx_full_path = r'D:\PythonWork\demos\outer_resource\docx\风险评估审查申请书.docx'
            pass
        dcmt = Document(docx_full_path)
        assert isinstance(dcmt, document.Document)
        # 需要编目录的 样式
        heading_styles = [
            'my_heading_1',
            'my_heading_2',
            'my_heading_3',
            'my_heading_4',
            'my_heading_5',
        ]
        heading_list = []
        for pgh in dcmt.paragraphs:
            if pgh.style.name in heading_styles:
                heading_text = pgh.text
                heading_lv = pgh.style.name
                my_heading_model = MyHeadingModel(text=heading_text, lv=heading_lv)
                heading_list.append(my_heading_model)
                pass  # end if..
            pass  # end for pgh..

        toc_flag = '__TOC__'
        for pgh in dcmt.paragraphs:
            if toc_flag in pgh.text:
                DocxComposer.__add_toc(dcmt, pgh, heading_list)
                break
                pass  # end if
            pass
        if not target_file:
            target_file = r'./1_toc_target.docx'
            pass
        dcmt.save(target_file)
        pass

    def unit_test(self):
        pass

    pass  # end of class DocxComposer


class DocxComposerApp:
    """应用系统检测报告/ 应用系统检测申请书 的 composer"""
    def __init__(self, session, declare_id, final_dcmt_path, template_type: TemplateType):
        self.session = session
        self.declare_id = declare_id
        self.final_dcmt_path = final_dcmt_path
        self.template_type = template_type

        if template_type == TemplateType.YingYong:
            self.BASE_DCMT = os.path.join(YINGYONG_DIR, r'chapters/_all_contents.docx')
        else:
            self.BASE_DCMT = os.path.join(YINGYONG_DECLARE_DIR, r'chapters/_all_contents.docx')
        pass

    def mailmerge_the_docx(self):
        """只有一个模板， 直接mailmerge"""
        # 直接mm
        from utils.docx_compose_mailmerge.test_mailmerge import MailMergeApp
        mm_app = MailMergeApp()
        mm_app.mailmerge_all(template_file=self.BASE_DCMT,
                             target_file=self.final_dcmt_path,
                             declare_id=self.declare_id,
                             session=self.session,
                             template_type=self.template_type)
        pass

    pass  # end of class DocxComposerApp


def __common_compose_api(template_type, session, report_id, final_dcmt_path):
    """公用的docx_compose_api"""
    # 创建随机的目录名称
    random_dir_name = globalIdGenerator.getNextId()
    docx_composer = DocxComposer(session=session,
                                 template_type=template_type,
                                 report_id=report_id,
                                 final_dcmt_path=final_dcmt_path,
                                 random_postfix=random_dir_name)
    docx_composer.final_compose()
    pass


def final_compose_evaluation_api(session, report_id, final_dcmt_path):
    """
    对测评文档 进行一键合并(compose) 以及 邮件合并(mailmerge)
    :param session: session
    :param report_id: report_id / declare_id
    :param final_dcmt_path:  最终生成的word路径
    :return:
    """
    __common_compose_api(template_type=TemplateType.CePing,
                         session=session,
                         report_id=report_id,
                         final_dcmt_path=final_dcmt_path)
    # 插入打分表
    from bll.evaluate.evaluation_task_bl import EvaluationTaskBll
    bl = EvaluationTaskBll()
    bl.export_result_table(declare_id=report_id, file_full_path=final_dcmt_path)
    pass


def final_compose_declare_api(session, report_id, final_dcmt_path):
    """
    对申报文档 进行一键合并(compose) 以及 邮件合并(mailmerge)

    :param session: session
    :param report_id: report_id / declare_id
    :param final_dcmt_path:  最终生成的word路径
    :return: None
    """
    __common_compose_api(template_type=TemplateType.ShenBao,
                         session=session,
                         report_id=report_id,
                         final_dcmt_path=final_dcmt_path)
    pass

def final_compose_fengping_api(session, report_id, final_dcmt_path):
    """
    对 风险评估报告 进行一键合并(compose) 以及 邮件合并(mailmerge)

    :param session: session
    :param report_id: report_id / declare_id
    :param final_dcmt_path:  最终生成的word路径
    :return: None
    """
    __common_compose_api(template_type=TemplateType.FengPing,
                         session=session,
                         report_id=report_id,
                         final_dcmt_path=final_dcmt_path)
    pass

def final_compose_fengping_declare_api(session, report_id, final_dcmt_path):
    """
    对 风险评估审查申请书 进行一键合并(compose) 以及 邮件合并(mailmerge)

    :param session: session
    :param report_id: report_id / declare_id
    :param final_dcmt_path:  最终生成的word路径
    :return: None
    """
    __common_compose_api(template_type=TemplateType.FengPingShenBao,
                         session=session,
                         report_id=report_id,
                         final_dcmt_path=final_dcmt_path)
    pass


def final_compose_yingYong_declare_api(session, declare_id, final_dcmt_path):
    """
    对 应用系统 审查申请书 进行邮件合并(mailmerge)

    :param session: session
    :param declare_id: report_id / declare_id
    :param final_dcmt_path:  最终生成的word路径
    :return: None
    """
    composer = DocxComposerApp(session=session,
                               declare_id=declare_id,
                               final_dcmt_path=final_dcmt_path,
                               template_type=TemplateType.YingYongShenBao)
    composer.mailmerge_the_docx()
    pass

def final_compose_yingYong_api(session, declare_id, final_dcmt_path):
    """
    对 应用系统检测报告 进行邮件合并(mailmerge)

    :param session: session
    :param declare_id: report_id / declare_id
    :param final_dcmt_path:  最终生成的word路径
    :return: None
    """
    # 也许要新开一个class， 因为应用系统检测报告， 是一份单独的内容
    composer = DocxComposerApp(session=session,
                               declare_id=declare_id,
                               final_dcmt_path=final_dcmt_path,
                               template_type=TemplateType.YingYong)
    composer.mailmerge_the_docx()
    pass


if __name__ == '__main__':
    # 完整测试需要的参数:session, report_id, final_target_dcmt
    DB_FILE_DECLARE = r'D:\PythonWork\DetectEvaluateService\declare.db'
    SessionFactory = DataBaseHelper.make_tmp_session_factory(complete_db_path=DB_FILE_DECLARE)
    the_session = SessionFactory()
    the_report_id = '1369945826411028480'
    final_compose_fengping_declare_api(session=the_session, report_id=the_report_id,
                                       final_dcmt_path='./final_result_fengping_declare.docx')
    pass
